#!/usr/bin/env python3

import mariadb
import docx
import docx2pdf


import os
import subprocess

import db_lib

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

docx_path = os.path.join(os.path.dirname(__file__), 'runmagedon.docx')
pdf_path = os.path.join(os.path.dirname(__file__), 'runmagedon.pdf')

def connect_to_db():
    login = db_lib.get_db_data()
    try: 
        conn = mariadb.connect(user = login['user'],
                            password = login['password'],
                            host = login['host'],
                            database = login['database'])
        
        cur = conn.cursor()
        return conn, cur 
    except mariadb.Error as e:
        print("Connection error:", e)
        return None, None
    
def get_names_turnament():
    conn, cursor = connect_to_db()
    if conn is None or cursor is None:
        print("❌ Cannot connect to the database.")
        return []

    query = """
        SELECT 
            r.kategoria,
            u.imie,
            u.nazwisko,
            r.wynik
        FROM runmagedon r
        JOIN uczestnik u ON r.id = u.runmagedon
        WHERE r.wynik IS NOT NULL
        ORDER BY r.kategoria, r.wynik;
    """

    # Use dictionary cursor for named access
    cursor = conn.cursor(dictionary=True)
    cursor.execute(query)
    rows = cursor.fetchall()
    cursor.close()
    conn.close()

    if not rows:
        print("⚠️ No race data available.")
        return []

    results_by_category = {}

    for row in rows:
        category = row['kategoria'] or "Brak kategorii"
        imie = row['imie']
        nazwisko = row['nazwisko']
        czas = row['wynik'] or "brak"
        results_by_category.setdefault(category, []).append([imie, nazwisko, czas])

    # Convert to requested list-of-dicts format
    formatted = [{category: entries} for category, entries in results_by_category.items()]
    return formatted
    
def generate_tournament_list(results, path=docx_path):
    """
    Create a .docx document with tournament results.
    Each group (category) is placed on a separate page,
    listing participants (imie, nazwisko, czas) sorted by shortest time.
    """
    document = docx.Document()

    for idx, group_data in enumerate(results):
        for category, participants in group_data.items():
            # --- Add category title ---
            title = document.add_paragraph(category if category else "Brak kategorii")
            for run in title.runs:
                run.bold = True
                run.font.size = Pt(20)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- Create table ---
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Imię'
            hdr_cells[1].text = 'Nazwisko'
            hdr_cells[2].text = 'Czas'

            # Make headers bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # --- Fill rows ---
            for imie, nazwisko, czas in participants:
                row_cells = table.add_row().cells
                row_cells[0].text = imie
                row_cells[1].text = nazwisko
                row_cells[2].text = czas

            # --- Add borders to the table ---
            tbl = table._tbl
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement('w:{}'.format(border_name))
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblPr = tbl.tblPr
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is None:
                    tblBorders = OxmlElement('w:tblBorders')
                    tblPr.append(tblBorders)
                tblBorders.append(border)

            # --- Page break after each category except the last one ---
            if idx < len(results) - 1:
                document.add_page_break()

    document.save(path)
    print(f"✅ Tournament list saved to: {path}")

def convert_docx_to_pdf(source_path=docx_path, destination_path=pdf_path):
    docx2pdf.convert(source_path, destination_path)
    
    
def convert_doc_to_pdf_linux(source_path=docx_path, destination_path=pdf_path):
    libreoffice_path = '/usr/bin/soffice'
    subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', source_path, '--outdir', os.path.dirname(destination_path)])
    print(f'Converted {source_path} to {destination_path}')
    
generate_tournament_list(get_names_turnament())
convert_doc_to_pdf_linux()
    
